"""
generate_docs.py — Builds the full backITup technical documentation as a .docx file.

Run standalone:
    python generate_docs.py

Auto-installs python-docx if it is missing, then writes
INTE403_VALILA_LHUISEGAHBRIELLE_BACKITUP_DOCUMENTATION.docx to the current directory.
"""

import subprocess
import sys

# ── Auto-install python-docx if missing ──────────────────────────────────────────
try:
    import docx  # noqa: F401
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call(
        [sys.executable, "-m", "pip", "install", "python-docx",
         "--break-system-packages", "--quiet"]
    )

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "INTE403_VALILA_LHUISEGAHBRIELLE_BACKITUP_DOCUMENTATION.docx"

BLACK = RGBColor(0, 0, 0)
FONT = "Times New Roman"
MONO = "Courier New"


# ── low-level font helper ─────────────────────────────────────────────────────────

def _set_run_font(run, name=FONT, size=12, bold=None, italic=None, color=BLACK):
    """Force a run's font family (incl. east-asian), size, weight, and color."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _style_body_paragraph(p, justify=True):
    """Apply the global body paragraph format: 1.15 spacing, 6pt after."""
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ── public helpers ────────────────────────────────────────────────────────────────

def add_heading(doc, text, level):
    """Add a heading with the required size/weight; level 1 is forced ALL CAPS."""
    sizes = {1: 14, 2: 13, 3: 12}
    if level == 1:
        text = text.upper()
    p = doc.add_heading("", level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_run_font(run, size=sizes.get(level, 12), bold=True, color=BLACK)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_body(doc, text, justify=True, italic=False, bold=False):
    """Add a justified 12pt Times New Roman body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=12, bold=bold, italic=italic, color=BLACK)
    _style_body_paragraph(p, justify=justify)
    return p


def add_rich_body(doc, segments, justify=True):
    """
    Add a body paragraph from a list of (text, bold) segments — used for inline
    bold labels (e.g. module names) within a normal paragraph.
    """
    p = doc.add_paragraph()
    for text, bold in segments:
        run = p.add_run(text)
        _set_run_font(run, size=12, bold=bold, color=BLACK)
    _style_body_paragraph(p, justify=justify)
    return p


def add_numbered(doc, items, justify=True):
    """Add a manually-numbered list (1. 2. 3. ...) as body paragraphs."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {item}")
        _set_run_font(run, size=12, color=BLACK)
        _style_body_paragraph(p, justify=justify)


def set_cell_black(cell, text, bold=False, mono=False, size=12, align_justify=False):
    """Write text into a table cell forcing black Times New Roman (or Courier New)."""
    cell.text = ""
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    if align_justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    _set_run_font(run, name=(MONO if mono else FONT), size=size, bold=bold, color=BLACK)


def add_table(doc, headers, rows, justify_last=False):
    """Add a bordered table with a bold header row, all text black."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        set_cell_black(table.rows[0].cells[j], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            justify = justify_last and j == len(row) - 1
            set_cell_black(cells[j], str(value), align_justify=justify)
    # spacing paragraph after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    return table


def add_monospace_block(doc, text):
    """Render a verbatim block in a 1-column table using Courier New 10pt."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    cell.text = ""
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        # keep the line literally (leading spaces preserved by python-docx)
        run = p.add_run(line if line != "" else " ")
        _set_run_font(run, name=MONO, size=10, color=BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    return table


def add_page_numbers(doc):
    """Right-aligned 'Page X' footer; the title (first) page shows no number."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True  # blank footer on page 1

    footer = section.footer  # applies to pages 2+
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run("Page ")
    _set_run_font(run, size=11, color=BLACK)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    field_run = p.add_run()
    _set_run_font(field_run, size=11, color=BLACK)
    field_run._r.append(fld_begin)
    field_run._r.append(instr)
    field_run._r.append(fld_end)


# ── document base style ───────────────────────────────────────────────────────────

def configure_base_styles(doc):
    """Make Normal (and thus inherited content) Times New Roman 12pt black."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)


def add_centered(doc, text, size=12, bold=False, italic=False):
    """Center-aligned line used for the title page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(2)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic, color=BLACK)
    return p


def blank_line(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    return p


# ══════════════════════════════════════════════════════════════════════════════════
#  DOCUMENT BUILD
# ══════════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()
    configure_base_styles(doc)
    add_page_numbers(doc)

    # ───────────────────────── TITLE PAGE ─────────────────────────
    blank_line(doc)
    blank_line(doc)
    add_centered(doc, "POLYTECHNIC UNIVERSITY OF THE PHILIPPINES", size=14, bold=True)
    add_centered(doc, "TAGUIG CAMPUS", size=14, bold=True)
    blank_line(doc)
    add_centered(doc, "COLLEGE OF INFORMATION AND COMMUNICATION TECHNOLOGY",
                 size=12, bold=True)
    blank_line(doc)
    blank_line(doc)
    add_centered(doc, "backITup", size=22, bold=True)
    add_centered(doc, "AUTOMATED BACKUP MANAGEMENT SYSTEM", size=16, bold=True)
    blank_line(doc)
    blank_line(doc)
    add_centered(doc, "A System Development Project", size=12)
    add_centered(doc, "Presented to the Faculty of the", size=12)
    add_centered(doc, "College of Information and Communication Technology", size=12)
    blank_line(doc)
    add_centered(doc, "In Partial Fulfillment of the Requirements for", size=12)
    add_centered(doc, "INTE403 – Systems Administration and Maintenance", size=12)
    blank_line(doc)
    blank_line(doc)
    add_centered(doc, "Submitted by:", size=12, bold=True)
    add_centered(doc, "VALILA, LHUISE GAHBRIELLE M.", size=12)
    add_centered(doc, "Section DIT 2-1", size=12)
    blank_line(doc)
    add_centered(doc, "Submitted to:", size=12, bold=True)
    add_centered(doc, "[Professor's Name]", size=12)
    add_centered(doc, "Course Instructor", size=12)
    blank_line(doc)
    blank_line(doc)
    add_centered(doc, "June 2026", size=12, bold=True)

    doc.add_page_break()

    # ───────────────────────── TABLE OF CONTENTS ─────────────────────────
    add_heading(doc, "Table of Contents", level=1)

    def leader(label, page, width=52):
        dots = "." * max(3, width - len(label) - len(str(page)))
        return f"{label} {dots} {page}"

    toc_entries = [
        ("Chapter 1 – Introduction", 3),
        ("Chapter 2 – System Architecture", 5),
        ("Chapter 3 – Installation Guide", 7),
        ("Chapter 4 – User Manual", 10),
        ("Chapter 5 – Security", 14),
        ("Chapter 6 – Performance & Recovery", 17),
        ("Chapter 7 – Email Notification System", 20),
        ("Annex A – Project File Structure", 22),
        ("Annex B – backup_systems.json Schema", 23),
        ("Annex C – users.json Schema", 24),
        ("Annex D – Glossary", 25),
    ]
    for label, page in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(leader(label, page))
        _set_run_font(run, name=MONO, size=11, color=BLACK)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15

    doc.add_page_break()

    # ═════════════════════ CHAPTER 1 – INTRODUCTION ═════════════════════
    add_heading(doc, "Chapter 1 – Introduction", level=1)

    add_heading(doc, "1.1 Background of the System", level=2)
    add_body(doc,
        "Reliable data backup has become a foundational requirement of modern "
        "information technology infrastructure. As individuals and organizations "
        "generate ever-larger volumes of digital files, the cost of losing that data "
        "to hardware failure, accidental deletion, ransomware, or simple human error "
        "has grown correspondingly severe. A dependable backup system is no longer a "
        "luxury reserved for large enterprises; it is a baseline operational safeguard "
        "that every computing environment should provide. Recognizing this, the "
        "discipline of systems administration places strong emphasis on designing "
        "backup and recovery strategies that are both automatic and verifiable.")
    add_body(doc,
        "Despite the clear need, most existing backup tools are poorly suited to "
        "students and small organizations. Enterprise-grade solutions tend to assume "
        "dedicated infrastructure, paid licensing, and administrators who are already "
        "fluent in their configuration. Lightweight consumer tools, on the other hand, "
        "often hide their inner workings behind opaque interfaces, making them "
        "unsuitable as a teaching vehicle. The result is a gap: there are few backup "
        "systems that are simultaneously simple to deploy, transparent in operation, "
        "and rich enough to demonstrate the full lifecycle of backup, encryption, "
        "monitoring, and recovery.")
    add_body(doc,
        "backITup was developed to fill exactly that gap. It is a Python-based backup "
        "management system targeting Ubuntu Linux that any user can launch with a "
        "single command, after which it automatically installs its own dependencies "
        "and guides the user through setup. Conceived as a learning project for the "
        "INTE403 Systems Administration and Maintenance course at the Polytechnic "
        "University of the Philippines – Taguig Campus, backITup deliberately exposes "
        "the concepts behind continuous synchronization, at-rest encryption, "
        "role-based access control, real-time monitoring, and recovery testing. In "
        "doing so it serves both as a practical utility and as a concrete study of how "
        "a complete backup solution is engineered.")

    add_heading(doc, "1.2 Objectives", level=2)
    add_body(doc,
        "The development of backITup was guided by a set of concrete objectives that "
        "together define the full scope of the system:")
    add_numbered(doc, [
        "Develop a continuous background backup daemon using Python's watchdog "
        "library to detect and mirror file changes in real time.",
        "Implement AES/Fernet encryption to secure all backed-up files at rest.",
        "Build a role-based authentication system using bcrypt password hashing.",
        "Create a real-time web dashboard using Flask and Socket.IO for monitoring.",
        "Provide email notifications for backup success and failure events via SMTP.",
        "Implement performance benchmarking with RTO/RPO tracking and restoration "
        "testing.",
        "Generate comprehensive technical documentation covering installation, "
        "operation, and recovery procedures.",
    ])

    add_heading(doc, "1.3 Scope and Limitations", level=2)
    add_body(doc,
        "The scope of backITup encompasses the complete backup lifecycle on a single "
        "Ubuntu Linux host running Python 3.8 or higher. The system provides "
        "continuous folder synchronization through a background daemon, optional "
        "per-system AES/Fernet encryption of all stored files, a bcrypt-backed "
        "authentication layer with administrator and user roles, and a real-time web "
        "dashboard for monitoring backup activity. It additionally delivers email "
        "notifications for backup outcomes over SMTP and a performance subsystem that "
        "benchmarks throughput, performs restoration tests, and computes Recovery Time "
        "and Recovery Point Objectives. All configuration and operation is handled "
        "through an interactive command-line menu and the accompanying dashboard.")
    add_body(doc,
        "Several limitations bound the present version of the system. backITup is "
        "Linux-only and offers no cross-platform support for Windows or macOS, owing "
        "to its reliance on Unix process semantics and the watchdog backend. It "
        "performs full mirror synchronization rather than incremental backup, and it "
        "does not de-duplicate encrypted files, so storage usage scales with the size "
        "of the watched data. The system does not provide cloud or off-site backup "
        "targets, nor does it ship a graphical installer; setup remains command-line "
        "driven. Finally, the bundled web dashboard runs on the Werkzeug development "
        "server and is therefore not hardened for production exposure, and should be "
        "placed behind a proper WSGI server and reverse proxy before any untrusted "
        "network use.")

    doc.add_page_break()

    # ═════════════════════ CHAPTER 2 – SYSTEM ARCHITECTURE ═════════════════════
    add_heading(doc, "Chapter 2 – System Architecture", level=1)

    add_heading(doc, "2.1 Architecture Overview", level=2)
    add_body(doc,
        "backITup follows a layered architecture that cleanly separates user "
        "interaction, business logic, supporting utilities, and persisted state. At "
        "the top sits the command-line entry point, backITup.py, which handles "
        "authentication, renders the role-aware menu, and dispatches user choices. "
        "Beneath it lies the core layer, a package of focused modules that implement "
        "configuration management, the backup engine, filesystem watching, "
        "scheduling, encryption, authentication, and benchmarking. A utilities layer "
        "provides cross-cutting helpers such as hashing, desktop notifications, and "
        "email alerts, while a data layer of JSON files and key material persists all "
        "configuration and results to disk.")
    add_body(doc,
        "Within the running system, three components cooperate to keep backups "
        "current. The watcher daemon, built on watchdog, observes the source folder "
        "and emits filesystem events the instant a file is created, modified, moved, "
        "or deleted. The scheduler independently triggers a complete synchronization "
        "at a configurable interval, guaranteeing convergence even if individual "
        "events are missed. Both feed into the backup engine, which performs the "
        "actual copy, optional encryption, and version rotation, and which emits live "
        "events and email alerts as work proceeds. This combination of event-driven "
        "and interval-driven backup provides both responsiveness and reliability.")

    add_heading(doc, "2.2 Module Breakdown", level=2)
    add_body(doc,
        "The system is composed of small, single-responsibility modules. The "
        "following describes each one and the role it plays within backITup:")
    modules = [
        ("core/config_manager.py", "manages backup system configurations stored in "
         "backup_systems.json, providing load, add, update, and remove operations."),
        ("core/backup_engine.py", "performs file copy and encryption, version "
         "rotation, and full synchronization, while emitting socket events and email "
         "alerts."),
        ("core/watcher.py", "a watchdog-based filesystem event listener that observes "
         "the source folder for changes."),
        ("core/event_handler.py", "maps raw filesystem events onto concrete backup "
         "engine actions such as back up, delete, or move."),
        ("core/scheduler.py", "runs a full sync on a configurable interval using the "
         "schedule library."),
        ("core/logger.py", "provides unified logging to both the console and "
         "logs/backITup.log."),
        ("core/encryption.py", "handles AES/Fernet key derivation via PBKDF2HMAC and "
         "the encryption and decryption of files."),
        ("core/auth.py", "implements user registration, bcrypt-based login, and role "
         "management."),
        ("core/benchmark.py", "provides performance benchmarking, restoration "
         "testing, and RTO/RPO calculation."),
        ("dashboard/app.py", "the Flask application factory wiring up the "
         "LoginManager and Socket.IO instance."),
        ("dashboard/routes.py", "defines the HTTP routes for all dashboard pages."),
        ("dashboard/socket_events.py", "broadcasts real-time Socket.IO events to "
         "connected browsers."),
        ("utils/file_utils.py", "supplies SHA256 hashing and path utilities."),
        ("utils/notify.py", "sends desktop notifications via plyer, failing silently "
         "when unavailable."),
        ("utils/email_alerts.py", "sends SMTP email alerts for backup success and "
         "failure."),
    ]
    for name, desc in modules:
        add_rich_body(doc, [(f"{name} ", True), ("— " + desc, False)])

    add_heading(doc, "2.3 Data Flow", level=2)
    add_body(doc,
        "When a file inside a watched folder changes, the flow begins at the watcher. "
        "The watchdog observer detects the filesystem event and passes it to the "
        "event handler, which translates the low-level event into a semantic action. "
        "That action is then invoked on the backup engine, which copies the file to "
        "the destination, encrypts it in place when encryption is enabled, and rotates "
        "older versions as configured. The same path is exercised by the scheduler's "
        "periodic full synchronization, ensuring the destination eventually mirrors "
        "the source even under heavy churn.")
    add_body(doc,
        "Live updates reach the web dashboard through a second, parallel flow. As the "
        "backup engine begins, progresses through, and completes a synchronization "
        "cycle, it calls into the socket events module, which emits named Socket.IO "
        "events such as backup_started, backup_progress, backup_complete, and "
        "backup_failed. The dashboard page, having opened a Socket.IO connection in "
        "the browser, receives these events and appends them to the live activity feed "
        "without any page refresh, giving the operator an immediate view of ongoing "
        "work.")
    add_body(doc,
        "Benchmark and recovery data follow a third flow centered on persistence. "
        "When a benchmark or restoration test is run, the benchmark module reads from "
        "the backup destination and the application log, computes its metrics, and "
        "prepends the result to data/benchmarks.json so that the newest entries appear "
        "first. The dashboard's Benchmarks tab then loads that same file and renders "
        "the backup and restore tables, attaching an RTO/RPO rating computed on the "
        "fly. In this way a single JSON file serves as the shared source of truth "
        "between the command-line tools and the web interface.")

    add_heading(doc, "2.4 Data Storage", level=2)
    add_body(doc,
        "All persistent state in backITup lives in plain files beneath the project "
        "directory, split between a data folder for configuration and results and a "
        "logs folder for textual output. The following table summarizes each file, "
        "where it is stored, and the purpose it serves:")
    add_table(doc,
        ["File", "Location", "Purpose"],
        [
            ["users.json", "data/", "User accounts and bcrypt hashes"],
            ["backup_systems.json", "data/", "Backup system configurations"],
            ["email_config.json", "data/", "SMTP email settings"],
            ["benchmarks.json", "data/", "Benchmark and restore test results"],
            ["dashboard_secret.key", "data/", "Flask session secret key"],
            ["backup.key", "data/<system_name>.key", "Per-system AES/Fernet key"],
            ["backITup.log", "logs/", "Unified application log"],
            ["benchmark_*.txt", "logs/", "Saved benchmark reports"],
        ])

    doc.add_page_break()

    # ═════════════════════ CHAPTER 3 – INSTALLATION GUIDE ═════════════════════
    add_heading(doc, "Chapter 3 – Installation Guide", level=1)

    add_heading(doc, "3.1 System Requirements", level=2)
    add_table(doc,
        ["Component", "Requirement"],
        [
            ["Operating System", "Ubuntu 22.04 LTS or later"],
            ["Python", "3.8 or higher"],
            ["Disk Space", "Minimum 100 MB (plus backup storage)"],
            ["RAM", "Minimum 256 MB"],
            ["Network", "Required for email alerts and dashboard remote access"],
            ["Privileges", "sudo access for installing python3-tk"],
        ])

    add_heading(doc, "3.2 Step-by-Step Installation", level=2)
    add_numbered(doc, [
        "Ensure Python 3.8 or higher is installed by running: python3 --version",
        "Clone or copy the backITup project folder to your machine: "
        "git clone https://github.com/yourrepo/backITup.git and then cd backITup",
        "No virtual environment is needed. Run the launcher directly: "
        "python backITup.py",
        "On first run, backITup auto-installs all dependencies: watchdog, schedule, "
        "plyer, cryptography, bcrypt, flask, flask-socketio, and flask-login.",
        "If tkinter is not installed, the system will prompt you to run: "
        "sudo apt-get install -y python3-tk",
        "After dependency installation the system restarts automatically and prompts "
        "you to create the first administrator account.",
    ])

    add_heading(doc, "3.3 First Run Walkthrough", level=2)
    add_body(doc,
        "The very first launch of backITup is primarily concerned with preparing the "
        "environment. When python backITup.py is executed, the launcher checks for "
        "each required third-party package and silently installs any that are missing "
        "before re-executing itself so the new modules are available. If the Tkinter "
        "bindings needed for the graphical folder picker are absent, it prints the "
        "single apt-get command required to install them. Once dependencies are "
        "satisfied and no user accounts yet exist, the program forces the creation of "
        "an administrator account, prompting for a username and a confirmed password "
        "that is immediately hashed with bcrypt and written to users.json.")
    add_body(doc,
        "With an administrator in place, the operator can create the first backup "
        "system by choosing Option 1 from the main menu. The system asks for a unique "
        "name and then opens a Tkinter folder picker to select the folder to watch, "
        "followed by a second picker for the backup destination. The operator is then "
        "asked whether to enable encryption and, if so, to enter and confirm an "
        "encryption password from which a Fernet key is derived and stored. Should the "
        "graphical picker be unavailable, the system gracefully falls back to manual "
        "path entry at the terminal.")
    add_body(doc,
        "As soon as a backup system is defined, its daemon begins working "
        "immediately. The backup engine performs an initial full synchronization to "
        "bring the destination into line with the source, the watcher starts "
        "listening for live filesystem events, and the scheduler arms the periodic "
        "full-sync interval. From this point forward the backup runs in the "
        "background, and the operator may safely return to the menu, launch the "
        "dashboard, or place the entire daemon into the background so it persists "
        "after the menu is closed.")

    add_heading(doc, "3.4 Running as a systemd Service", level=2)
    add_body(doc,
        "For unattended operation it is desirable to run backITup as a systemd "
        "service so that the backup daemon survives reboots and is automatically "
        "restarted if it ever fails. The following unit file, saved as "
        "/etc/systemd/system/backitup.service, defines such a service; replace "
        "YOUR_USERNAME and the project path with the values appropriate to your "
        "installation:")
    add_monospace_block(doc,
        "[Unit]\n"
        "Description=backITup Automated Backup Daemon\n"
        "After=network.target\n"
        "\n"
        "[Service]\n"
        "Type=simple\n"
        "User=YOUR_USERNAME\n"
        "WorkingDirectory=/path/to/backITup\n"
        "ExecStart=/usr/bin/python3 /path/to/backITup/backITup.py\n"
        "Restart=on-failure\n"
        "RestartSec=10\n"
        "\n"
        "[Install]\n"
        "WantedBy=multi-user.target")
    add_body(doc,
        "Once the unit file is in place, enable and start the service with the "
        "following commands:")
    add_monospace_block(doc,
        "sudo systemctl daemon-reexec\n"
        "sudo systemctl enable backitup\n"
        "sudo systemctl start backitup\n"
        "sudo systemctl status backitup")

    doc.add_page_break()

    # ═════════════════════ CHAPTER 4 – USER MANUAL ═════════════════════
    add_heading(doc, "Chapter 4 – User Manual", level=1)

    add_heading(doc, "4.1 CLI Main Menu Overview", level=2)
    add_body(doc,
        "After logging in, the operator is presented with a numbered text menu whose "
        "available options depend on the account's role. Administrators see the full "
        "set of management and recovery functions, while standard users see only the "
        "options for viewing status and triggering a manual backup. The table below "
        "lists the administrator menu and the action each option performs:")
    add_table(doc,
        ["Option", "Description"],
        [
            ["[1]", "Create a new backup system"],
            ["[2]", "Configure a current backup system"],
            ["[3]", "Delete a backup system"],
            ["[4]", "View backup status"],
            ["[5]", "Trigger manual backup"],
            ["[6]", "Manage users (Admin only)"],
            ["[7]", "Run in background"],
            ["[8]", "Launch Web Dashboard (Admin only)"],
            ["[9]", "Performance & Recovery (Admin only)"],
            ["[10]", "Logout"],
            ["[11/12]", "Exit"],
        ])

    add_heading(doc, "4.2 Creating a Backup System", level=2)
    add_body(doc,
        "Creating a backup system is a short guided sequence carried out from the "
        "main menu:")
    add_numbered(doc, [
        "Select Option [1] from the main menu to begin creating a backup system.",
        "Enter a unique name for the backup system when prompted.",
        "Use the Tkinter folder picker to select the watch folder — the source "
        "directory that will be monitored.",
        "Use the Tkinter folder picker again to select the backup destination where "
        "copies will be stored.",
        "Choose whether to enable encryption by answering Y or N.",
        "If encryption is enabled, enter and then confirm an encryption password; a "
        "Fernet key is derived from it and saved.",
        "The system writes the configuration, saves the key, performs an initial full "
        "sync, and starts the daemon in the background.",
    ])

    add_heading(doc, "4.3 Configuring a Backup System", level=2)
    add_body(doc,
        "An existing backup system can be adjusted at any time through Option [2] on "
        "the main menu. The configuration flow allows the operator to change the "
        "backup destination folder, the interval in minutes at which the periodic full "
        "synchronization runs, and the maximum number of previous file versions to "
        "retain. The same Configure menu also provides access to the global email "
        "alert settings, allowing SMTP notifications to be set up or updated without "
        "leaving the backup configuration workflow.")
    add_body(doc,
        "One property is deliberately immutable: the source watch folder is locked "
        "once a backup system has been created. Because the entire backup history, "
        "version chain, and — when encryption is enabled — the relationship between "
        "source files and their encrypted copies are all anchored to the original "
        "source path, permitting it to change would risk orphaning existing backups "
        "and breaking restoration. To point backups at a different source, the "
        "operator instead deletes the system and creates a new one, preserving the "
        "integrity of every existing configuration.")

    add_heading(doc, "4.4 Web Dashboard Guide", level=2)

    add_heading(doc, "4.4.1 Login Page", level=3)
    add_body(doc,
        "The dashboard is reached in a browser at http://localhost:5000. It accepts "
        "the same credentials used for the command-line login, since both read from "
        "the shared users.json account store. A failed login does not reveal which "
        "field was incorrect; it simply redisplays the form with a red error message "
        "indicating that the username or password was invalid.")

    add_heading(doc, "4.4.2 Dashboard (Main Page)", level=3)
    add_body(doc,
        "The main page presents every configured backup system as a card. Each card "
        "displays the system name, the source and destination paths, the encryption "
        "status, the time of the most recent backup, and a colored status badge. A "
        "Live Activity Feed beneath the cards updates in real time over Socket.IO, "
        "appending a new line each time a backup starts, progresses, completes, or "
        "fails — all without any page refresh. Administrators additionally see Backup "
        "Now and Stop/Resume controls on each card for direct, per-system action.")

    add_heading(doc, "4.4.3 Job History Page", level=3)
    add_body(doc,
        "The Job History page renders a table of the last one hundred backup events "
        "parsed directly from backITup.log, newest first. A filter bar lets the "
        "operator narrow the view by date and by status, choosing among All, Success, "
        "and Failed. A second tab on the same page, labeled Benchmarks, presents the "
        "recorded backup benchmark results alongside restoration test results, each "
        "restore row annotated with its RTO/RPO rating badge.")

    add_heading(doc, "4.4.4 User Management Page (Admin only)", level=3)
    add_body(doc,
        "Available only to administrators, the User Management page lists every "
        "account together with a role badge distinguishing admins from standard "
        "users. An Add User form beneath the list accepts a username, a password, and "
        "a role, creating the account with a freshly bcrypt-hashed password. Each "
        "existing user carries a Delete button guarded by a confirmation dialog, and "
        "the interface prevents administrators from deleting the very account they are "
        "currently logged in with.")

    add_heading(doc, "4.5 Email Alert Configuration", level=2)
    add_body(doc,
        "Email alerts are configured through the Configure menu and verified "
        "immediately with a live test message:")
    add_numbered(doc, [
        "Select Option [2] Configure from the main menu.",
        "Choose [2] Email alerts from the configuration sub-menu.",
        "Enter the SMTP host, for example smtp.gmail.com.",
        "Enter the SMTP port; the default is 587.",
        "Enter the sender Gmail address.",
        "Enter the Gmail App Password — not your ordinary Gmail login password.",
        "Enter the recipient email address.",
        "A test email is sent immediately to confirm that the configuration works.",
    ])
    add_body(doc,
        "Note: For Gmail, users must generate an App Password under Google Account > "
        "Security > 2-Step Verification > App Passwords. The regular Gmail password "
        "will not work with SMTP, because Google requires application-specific "
        "credentials for programmatic mail access once two-step verification is "
        "enabled.")

    doc.add_page_break()

    # ═════════════════════ CHAPTER 5 – SECURITY ═════════════════════
    add_heading(doc, "Chapter 5 – Security", level=1)

    add_heading(doc, "5.1 Encryption Implementation", level=2)
    add_body(doc,
        "backITup protects backed-up data at rest using symmetric encryption provided "
        "by the Python cryptography library's Fernet recipe, which is built on "
        "AES-128 in CBC mode authenticated with HMAC-SHA256. Every encrypted file is "
        "therefore both confidential and tamper-evident: any modification to the "
        "stored ciphertext is detected when the authentication tag fails to verify "
        "during decryption. Encryption is configured independently for each backup "
        "system, so an operator may keep some sources in plain mirrored form while "
        "securing others.")
    add_body(doc,
        "The encryption key is never stored directly from the user's password. "
        "Instead, backITup derives it with PBKDF2HMAC using SHA256 as the underlying "
        "hash, 480,000 iterations, and a 16-byte random salt. This deliberately "
        "expensive derivation frustrates brute-force attacks against the password, "
        "while the per-system salt ensures that two systems protected by the same "
        "password still produce entirely different keys. The salt is stored in "
        "Base64 form inside backup_systems.json alongside the rest of the system's "
        "configuration so that the identical key can be reconstructed when needed.")
    add_body(doc,
        "Each system's derived key is written to its own file at data/<name>.key, and "
        "the configuration records the path to that key. During a backup, files are "
        "encrypted directly from the source into the destination and given a .enc "
        "suffix; the original source files are never altered in the process. This "
        "design keeps the plaintext untouched on the source side while ensuring that "
        "everything written to the backup destination is encrypted, and it allows the "
        "restoration subsystem to decrypt a sampled file and verify it byte-for-byte "
        "against its untouched source.")

    add_heading(doc, "5.2 Authentication and Role-Based Access Control", level=2)
    add_body(doc,
        "User authentication is built on bcrypt, an adaptive password-hashing "
        "function. When an account is created, the password is combined with a random "
        "salt and processed through bcrypt's cost-factor-controlled key schedule to "
        "produce a single self-describing hash that embeds both the salt and the cost "
        "factor. Only this hash is stored in users.json; the plaintext password is "
        "never written to disk. At login, the supplied password is verified with "
        "bcrypt's constant-time checkpw routine, which re-derives the hash and "
        "compares it safely against the stored value.")
    add_body(doc,
        "Access is governed by a two-tier role model. An administrator may create, "
        "configure, and delete backup systems, manage user accounts, launch the web "
        "dashboard, and run performance benchmarks and recovery tests. A standard "
        "user is restricted to viewing backup status and triggering a manual backup. "
        "The following matrix summarizes which capabilities are available to each "
        "role:")
    add_table(doc,
        ["Feature", "Admin", "User"],
        [
            ["Create backup system", "Yes", "No"],
            ["Configure backup system", "Yes", "No"],
            ["Delete backup system", "Yes", "No"],
            ["View backup status", "Yes", "Yes"],
            ["Trigger manual backup", "Yes", "Yes"],
            ["Manage users", "Yes", "No"],
            ["Launch web dashboard", "Yes", "No"],
            ["Performance & Recovery", "Yes", "No"],
        ])

    add_heading(doc, "5.3 Access Restrictions", level=2)
    add_body(doc,
        "On the command line, the operator's role is established at login and checked "
        "before every privileged action. When a standard user attempts to reach an "
        "administrator-only option, the system does not perform the action; instead it "
        "prints the message \"Permission denied — admin only\" and returns cleanly to "
        "the menu. Because the check is performed at the point of use rather than "
        "merely hidden in the menu rendering, the restriction holds even if a user "
        "enters an option number directly.")
    add_body(doc,
        "On the web side, Flask-Login guards every route so that unauthenticated "
        "requests are redirected to the login page. Administrator-only routes apply an "
        "additional role check and respond to non-administrators with an HTTP 403 "
        "status and a plain HTML message explaining that administrator privileges are "
        "required. Browser sessions are signed with a cryptographically random "
        "24-byte secret key that is generated once and stored in "
        "data/dashboard_secret.key, ensuring that session cookies cannot be forged "
        "across restarts.")

    doc.add_page_break()

    # ═════════════════════ CHAPTER 6 – PERFORMANCE & RECOVERY ═════════════════════
    add_heading(doc, "Chapter 6 – Performance & Recovery", level=1)

    add_heading(doc, "6.1 Benchmarking Overview", level=2)
    add_body(doc,
        "The benchmarking subsystem measures how much data a backup system is "
        "protecting and how quickly that data can be traversed. For a chosen system "
        "it walks the entire backup destination, counting the total number of files "
        "and summing their sizes, and times the operation with a high-resolution "
        "clock. From these measurements it computes the total size in megabytes, the "
        "duration in seconds, and the resulting throughput in megabytes per second, "
        "giving a concrete picture of the system's scale and the speed of the "
        "underlying storage.")
    add_body(doc,
        "Every benchmark result is stored as an entry in data/benchmarks.json, "
        "prepended so that the most recent run appears first. The results can be "
        "reviewed in two ways: through Option [9] Performance & Recovery on the "
        "command-line menu, which prints the figures directly to the terminal, or "
        "through the Benchmarks tab of the web dashboard, which renders them in a "
        "readable table alongside the restoration test history.")

    add_heading(doc, "6.2 Restoration Testing", level=2)
    add_body(doc,
        "Restoration testing verifies that backups are not merely present but "
        "genuinely recoverable. The subsystem selects up to three files at random "
        "from the backup destination and restores each one into a temporary working "
        "directory. When the system is encrypted, the sampled file is decrypted with "
        "the system's key during restoration; otherwise it is copied directly. The "
        "restored file's SHA256 hash is then compared against the hash of the "
        "corresponding original source file, and the test passes only if the two "
        "digests match exactly, proving the backup is a faithful reproduction.")
    add_body(doc,
        "Robustness is built into the procedure. The temporary directory used for "
        "restoration is always removed in a finally block, so it is cleaned up "
        "regardless of whether any individual file restore succeeds, fails, or raises "
        "an exception. Each file's pass or fail outcome, along with its restore time, "
        "is recorded individually, and the aggregate result — counts of passed and "
        "failed files and the total restore time — is saved to benchmarks.json for "
        "later analysis and for the RTO calculation described below.")

    add_heading(doc, "6.3 RTO/RPO Methodology", level=2)
    add_body(doc,
        "The Recovery Time Objective expresses how long it takes to recover data. "
        "backITup estimates it as the average total restore time across the last five "
        "restoration tests for a system, measured in seconds using Python's "
        "high-resolution time.perf_counter clock. Because the figure is derived from "
        "actual restorations of real backed-up files, it reflects the true cost of "
        "recovery on the host's hardware rather than a theoretical estimate.")
    add_body(doc,
        "The Recovery Point Objective expresses how much recent data might be lost "
        "between backups. backITup computes it by parsing backITup.log for successful "
        "full-sync events for the system, taking the two most recent timestamps and "
        "reporting the gap between them in minutes. This represents the interval "
        "during which newly changed data would not yet have been captured. The "
        "measured RTO and RPO are then mapped to a qualitative rating using the "
        "thresholds in the table below:")
    add_table(doc,
        ["Rating", "RTO Threshold", "RPO Threshold"],
        [
            ["Excellent", "< 30 seconds", "< 5 minutes"],
            ["Good", "< 120 seconds", "< 30 minutes"],
            ["Poor", ">= 120 seconds", ">= 30 minutes"],
            ["Insufficient Data", "Not enough data", "Not enough data"],
        ])

    add_heading(doc, "6.4 Sample Benchmark Report", level=2)
    add_body(doc,
        "Choosing Generate Full Report produces a plain-text performance and recovery "
        "summary that is both printed to the terminal and saved under the logs folder. "
        "A representative report is shown below:")
    add_monospace_block(doc,
        "=======================================================\n"
        "  backITup - Performance & Recovery Report\n"
        "  System : Documents Backup\n"
        "  Date   : 2026-06-25 14:00:00\n"
        "=======================================================\n"
        "\n"
        "-- Backup Benchmarks (last 5) -------------------------\n"
        "  Timestamp              Files   Size (MB)  Duration (s)  Throughput (MB/s)\n"
        "  2026-06-25 13:55:00      142       18.340        0.8821            20.791\n"
        "  2026-06-25 13:25:00      141       18.210        0.9102            19.896\n"
        "  2026-06-25 12:55:00      139       17.980        0.8843            20.331\n"
        "\n"
        "-- Restore Tests (last 5) ----------------------------\n"
        "  Timestamp              Files Tested  Passed  Failed  Total Time (s)\n"
        "  2026-06-25 13:58:00               3       3       0          0.0082\n"
        "  2026-06-25 13:28:00               3       3       0          0.0094\n"
        "\n"
        "-- RTO / RPO Summary ---------------------------------\n"
        "  Average RTO : 0.0088s  (target: < 30s for Excellent)\n"
        "  RPO         : 30.0 min (target: < 5 min for Excellent)\n"
        "  Rating      : Good\n"
        "\n"
        "=======================================================")

    doc.add_page_break()

    # ═════════════════════ CHAPTER 7 – EMAIL NOTIFICATION SYSTEM ═════════════════════
    add_heading(doc, "Chapter 7 – Email Notification System", level=1)

    add_heading(doc, "7.1 SMTP Configuration", level=2)
    add_body(doc,
        "backITup sends its notifications using only Python's built-in smtplib, "
        "establishing a connection to the configured mail server and upgrading it to "
        "an encrypted channel with STARTTLS on the standard submission port 587. "
        "Because it relies on no third-party email libraries, the notification system "
        "remains lightweight and portable, working with any standards-compliant SMTP "
        "provider such as Gmail, Outlook, or a private mail relay.")
    add_body(doc,
        "All mail settings are persisted in data/email_config.json, including the SMTP "
        "host and port, the sender address and application password, the recipient "
        "address, and an enabled flag. That flag is central to safe operation: until "
        "the configuration has been completed and the flag set, the system refuses to "
        "attempt delivery, which prevents half-configured or empty settings from "
        "generating spurious connection errors during routine backups.")

    add_heading(doc, "7.2 Alert Types", level=2)
    add_body(doc,
        "backITup distinguishes between two kinds of operational alert, each with its "
        "own subject line and color-themed body so that the outcome of a backup is "
        "obvious at a glance:")
    add_rich_body(doc, [
        ("Success Alert: ", True),
        ("the subject line combines a check-mark emoji with \"backITup — Backup "
         "Complete: {system_name}\". The body presents the system name, the count of "
         "files backed up, and the duration in seconds, capped by a green SUCCESS "
         "status badge.", False),
    ])
    add_rich_body(doc, [
        ("Failure Alert: ", True),
        ("the subject line combines a warning emoji with \"backITup — Backup FAILED: "
         "{system_name}\". The body presents the system name, the error message, and "
         "the timestamp of the failure, capped by a red FAILED status badge.", False),
    ])

    add_heading(doc, "7.3 Email Template Design", level=2)
    add_body(doc,
        "The notification emails are formatted as HTML using inline CSS exclusively, "
        "with no external stylesheets. This is a deliberate choice, since mail clients "
        "routinely strip linked stylesheets; inlining every style rule guarantees that "
        "the message renders consistently across webmail and desktop clients alike. "
        "The visual design centers on a dark header bar in the shade #1e1e2e carrying "
        "the backITup branding, above a white content area constrained to a maximum "
        "width of 600 pixels, a detail table with alternating row colors for "
        "legibility, and a small footer reading \"Sent by backITup Automated Backup "
        "System\".")
    add_body(doc,
        "Equally important is the system's failure behavior. The alert routines are "
        "written to fail silently: if email has not been configured, or if any error "
        "occurs while connecting or sending, the function records a warning or error "
        "in the application log and returns without raising an exception. This ensures "
        "that the email subsystem can never interrupt or abort a backup; notification "
        "is treated as a best-effort convenience layered on top of the core "
        "protection the system provides.")

    doc.add_page_break()

    # ═════════════════════ ANNEX A – PROJECT FILE STRUCTURE ═════════════════════
    add_heading(doc, "Annex A – Project File Structure", level=1)
    add_body(doc,
        "The complete directory layout of the backITup project is shown below. The "
        "tree reflects the layered architecture described in Chapter 2, with the "
        "command-line entry point at the root and the core, dashboard, utils, data, "
        "and logs directories beneath it:")
    add_monospace_block(doc,
        "backITup/\n"
        "├── backITup.py\n"
        "├── core/\n"
        "│   ├── __init__.py\n"
        "│   ├── auth.py\n"
        "│   ├── backup_engine.py\n"
        "│   ├── benchmark.py\n"
        "│   ├── config_manager.py\n"
        "│   ├── encryption.py\n"
        "│   ├── event_handler.py\n"
        "│   ├── logger.py\n"
        "│   ├── scheduler.py\n"
        "│   └── watcher.py\n"
        "├── dashboard/\n"
        "│   ├── __init__.py\n"
        "│   ├── app.py\n"
        "│   ├── routes.py\n"
        "│   ├── socket_events.py\n"
        "│   └── templates/\n"
        "│       ├── base.html\n"
        "│       ├── index.html\n"
        "│       ├── jobs.html\n"
        "│       ├── login.html\n"
        "│       └── users.html\n"
        "├── utils/\n"
        "│   ├── __init__.py\n"
        "│   ├── email_alerts.py\n"
        "│   ├── file_utils.py\n"
        "│   └── notify.py\n"
        "├── data/\n"
        "│   ├── backup_systems.json\n"
        "│   ├── benchmarks.json\n"
        "│   ├── email_config.json\n"
        "│   └── users.json\n"
        "└── logs/\n"
        "    ├── backITup.log\n"
        "    └── benchmark_*.txt")

    doc.add_page_break()

    # ═════════════════════ ANNEX B – backup_systems.json SCHEMA ═════════════════════
    add_heading(doc, "Annex B – backup_systems.json Schema", level=1)
    add_body(doc,
        "Each backup system is stored as a JSON object within the array held in "
        "backup_systems.json. The fields that make up a single system configuration "
        "are described in the following table:")
    add_table(doc,
        ["Field", "Type", "Description"],
        [
            ["name", "string", "Unique name for the backup system"],
            ["watch_folder", "string", "Absolute path of the folder being monitored"],
            ["backup_destination", "string", "Absolute path where backups are stored"],
            ["interval_minutes", "integer", "How often a full sync runs (in minutes)"],
            ["max_versions", "integer", "Number of previous file versions to retain"],
            ["encryption_enabled", "boolean", "Whether AES/Fernet encryption is active"],
            ["salt", "string", "Base64-encoded PBKDF2 salt for key derivation"],
            ["key_path", "string", "Absolute path to the saved Fernet key file"],
        ])

    doc.add_page_break()

    # ═════════════════════ ANNEX C – users.json SCHEMA ═════════════════════
    add_heading(doc, "Annex C – users.json Schema", level=1)
    add_body(doc,
        "User accounts are stored as JSON objects within the array held in "
        "users.json. Each account is described by the following fields:")
    add_table(doc,
        ["Field", "Type", "Description"],
        [
            ["username", "string", "Unique login identifier"],
            ["password_hash", "string", "bcrypt hash of the user's password"],
            ["role", "string", "Either \"admin\" or \"user\""],
            ["created_at", "string", "ISO-format datetime of account creation"],
        ])

    doc.add_page_break()

    # ═════════════════════ ANNEX D – GLOSSARY ═════════════════════
    add_heading(doc, "Annex D – Glossary", level=1)
    add_body(doc,
        "The following glossary defines the principal technical terms used throughout "
        "this documentation:")
    add_table(doc,
        ["Term", "Definition"],
        [
            ["AES", "Advanced Encryption Standard; symmetric block cipher used via "
             "Fernet"],
            ["bcrypt", "Adaptive password hashing function with built-in salting and "
             "cost factor"],
            ["Daemon", "A background process that runs continuously without direct "
             "user interaction"],
            ["Fernet", "A symmetric encryption scheme built on AES-128-CBC + "
             "HMAC-SHA256"],
            ["Flask", "A lightweight Python web framework used for the backITup "
             "dashboard"],
            ["PBKDF2HMAC", "Password-Based Key Derivation Function 2; used to derive "
             "encryption keys from passwords"],
            ["RPO", "Recovery Point Objective; the maximum acceptable age of backup "
             "data"],
            ["RTO", "Recovery Time Objective; the maximum acceptable time to restore "
             "a backup"],
            ["Salt", "A random value added to a password before hashing to prevent "
             "rainbow table attacks"],
            ["SHA256", "A cryptographic hash function producing a 256-bit digest; "
             "used for file integrity checks"],
            ["Socket.IO", "A library for real-time bidirectional communication "
             "between client and server"],
            ["SMTP", "Simple Mail Transfer Protocol; used to send email "
             "notifications"],
            ["STARTTLS", "An extension to SMTP that upgrades a plain connection to "
             "TLS encryption"],
            ["Watcher", "A filesystem monitor (via watchdog) that detects file "
             "changes in real time"],
            ["watchdog", "Python library for monitoring filesystem events"],
        ],
        justify_last=True)

    return doc


def main():
    doc = build_document()
    doc.save(OUTPUT_FILE)
    print(f"Documentation saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
